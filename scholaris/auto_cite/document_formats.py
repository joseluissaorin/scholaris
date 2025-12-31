"""Multi-format document input/output support for auto-citation system.

Supports reading and writing documents in multiple formats:
- Plain text (.txt)
- Markdown (.md)
- Microsoft Word (.docx)
- PDF (.pdf) - input only
- HTML (.html)
- Rich Text Format (.rtf)
- OpenDocument Text (.odt)
- LaTeX (.tex)

Usage:
    # Input
    text = DocumentFormatHandler.read_document("paper.docx")

    # Output
    DocumentFormatHandler.write_document(cited_text, "output.docx", style="apa")

    # Auto-detect format
    text = DocumentFormatHandler.read_document_auto("input.pdf")
"""

import os
import re
from pathlib import Path
from typing import Optional, Dict, Any, Literal
from enum import Enum

# Core dependencies
try:
    from docx import Document
    from docx.shared import Pt, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

try:
    from bs4 import BeautifulSoup
    import markdown
    HTML_AVAILABLE = True
except ImportError:
    HTML_AVAILABLE = False

try:
    import pypandoc
    PANDOC_AVAILABLE = True
except ImportError:
    PANDOC_AVAILABLE = False

try:
    from odf import text, teletype
    from odf.opendocument import load as odf_load, OpenDocumentText
    ODT_AVAILABLE = True
except ImportError:
    ODT_AVAILABLE = False

try:
    from pypdf import PdfReader
    PDF_AVAILABLE = True
except ImportError:
    try:
        from PyPDF2 import PdfReader
        PDF_AVAILABLE = True
    except ImportError:
        PDF_AVAILABLE = False

try:
    from weasyprint import HTML as WeasyHTML
    WEASYPRINT_AVAILABLE = True
except ImportError:
    WEASYPRINT_AVAILABLE = False

try:
    from reportlab.lib.pagesizes import letter
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
    from reportlab.lib.styles import getSampleStyleSheet
    REPORTLAB_AVAILABLE = True
except ImportError:
    REPORTLAB_AVAILABLE = False


class DocumentFormat(str, Enum):
    """Supported document formats."""
    TXT = "txt"
    MARKDOWN = "md"
    DOCX = "docx"
    PDF = "pdf"
    HTML = "html"
    RTF = "rtf"
    ODT = "odt"
    LATEX = "tex"


class DocumentFormatHandler:
    """Handles reading and writing documents in multiple formats."""

    # Format to extension mapping
    FORMAT_EXTENSIONS = {
        DocumentFormat.TXT: [".txt", ".text"],
        DocumentFormat.MARKDOWN: [".md", ".markdown"],
        DocumentFormat.DOCX: [".docx"],
        DocumentFormat.PDF: [".pdf"],
        DocumentFormat.HTML: [".html", ".htm"],
        DocumentFormat.RTF: [".rtf"],
        DocumentFormat.ODT: [".odt"],
        DocumentFormat.LATEX: [".tex", ".latex"],
    }

    @classmethod
    def detect_format(cls, file_path: str) -> DocumentFormat:
        """Auto-detect document format from file extension.

        Args:
            file_path: Path to the document

        Returns:
            Detected DocumentFormat

        Raises:
            ValueError: If format cannot be detected
        """
        ext = Path(file_path).suffix.lower()

        for fmt, extensions in cls.FORMAT_EXTENSIONS.items():
            if ext in extensions:
                return fmt

        raise ValueError(
            f"Unsupported file extension: {ext}. "
            f"Supported: {', '.join(sum(cls.FORMAT_EXTENSIONS.values(), []))}"
        )

    @classmethod
    def read_document(
        cls,
        file_path: str,
        format: Optional[DocumentFormat] = None
    ) -> str:
        """Read document from file in specified format.

        Args:
            file_path: Path to the document
            format: Document format (auto-detected if None)

        Returns:
            Document text content

        Raises:
            ValueError: If format is not supported
            FileNotFoundError: If file doesn't exist
        """
        if not os.path.exists(file_path):
            raise FileNotFoundError(f"Document not found: {file_path}")

        # Auto-detect format if not specified
        if format is None:
            format = cls.detect_format(file_path)

        # Route to appropriate reader
        readers = {
            DocumentFormat.TXT: cls._read_txt,
            DocumentFormat.MARKDOWN: cls._read_markdown,
            DocumentFormat.DOCX: cls._read_docx,
            DocumentFormat.PDF: cls._read_pdf,
            DocumentFormat.HTML: cls._read_html,
            DocumentFormat.RTF: cls._read_rtf,
            DocumentFormat.ODT: cls._read_odt,
            DocumentFormat.LATEX: cls._read_latex,
        }

        reader = readers.get(format)
        if reader is None:
            raise ValueError(f"Unsupported input format: {format}")

        return reader(file_path)

    @classmethod
    def write_document(
        cls,
        text: str,
        output_path: str,
        format: Optional[DocumentFormat] = None,
        citation_style: Literal["apa", "chicago"] = "apa",
        metadata: Optional[Dict[str, Any]] = None
    ) -> None:
        """Write document to file in specified format.

        Args:
            text: Document text content (with citations)
            output_path: Path to save the document
            format: Document format (auto-detected if None)
            citation_style: Citation style for formatting
            metadata: Optional metadata (title, author, etc.)

        Raises:
            ValueError: If format is not supported
        """
        # Auto-detect format if not specified
        if format is None:
            format = cls.detect_format(output_path)

        metadata = metadata or {}

        # Route to appropriate writer
        writers = {
            DocumentFormat.TXT: cls._write_txt,
            DocumentFormat.MARKDOWN: cls._write_markdown,
            DocumentFormat.DOCX: cls._write_docx,
            DocumentFormat.HTML: cls._write_html,
            DocumentFormat.RTF: cls._write_rtf,
            DocumentFormat.ODT: cls._write_odt,
            DocumentFormat.LATEX: cls._write_latex,
            DocumentFormat.PDF: cls._write_pdf,
        }

        writer = writers.get(format)
        if writer is None:
            raise ValueError(f"Unsupported output format: {format}")

        writer(text, output_path, citation_style, metadata)

    # ========== INPUT READERS ==========

    @staticmethod
    def _read_txt(file_path: str) -> str:
        """Read plain text file."""
        with open(file_path, 'r', encoding='utf-8') as f:
            return f.read()

    @staticmethod
    def _read_markdown(file_path: str) -> str:
        """Read Markdown file."""
        with open(file_path, 'r', encoding='utf-8') as f:
            return f.read()

    @staticmethod
    def _read_docx(file_path: str) -> str:
        """Read Microsoft Word document."""
        if not DOCX_AVAILABLE:
            raise ImportError("python-docx is required for DOCX support. Install: pip install python-docx")

        doc = Document(file_path)
        paragraphs = [p.text for p in doc.paragraphs]
        return '\n\n'.join(paragraphs)

    @staticmethod
    def _read_pdf(file_path: str) -> str:
        """Read PDF document (text extraction)."""
        if not PDF_AVAILABLE:
            raise ImportError("pypdf or PyPDF2 is required for PDF support. Install: pip install pypdf")

        reader = PdfReader(file_path)
        text_parts = []
        for page in reader.pages:
            text = page.extract_text()
            if text:
                text_parts.append(text)

        return '\n\n'.join(text_parts)

    @staticmethod
    def _read_html(file_path: str) -> str:
        """Read HTML document."""
        if not HTML_AVAILABLE:
            raise ImportError("beautifulsoup4 is required for HTML support. Install: pip install beautifulsoup4")

        with open(file_path, 'r', encoding='utf-8') as f:
            html_content = f.read()

        soup = BeautifulSoup(html_content, 'html.parser')

        # Remove script and style elements
        for script in soup(["script", "style"]):
            script.extract()

        # Get text
        text = soup.get_text()

        # Clean up whitespace
        lines = (line.strip() for line in text.splitlines())
        chunks = (phrase.strip() for line in lines for phrase in line.split("  "))
        text = '\n'.join(chunk for chunk in chunks if chunk)

        return text

    @staticmethod
    def _read_rtf(file_path: str) -> str:
        """Read RTF document using pandoc."""
        if PANDOC_AVAILABLE:
            # Use pandoc if available
            return pypandoc.convert_file(file_path, 'plain', format='rtf')
        else:
            raise ImportError(
                "pypandoc (with pandoc installed) is required for RTF support. "
                "Install: pip install pypandoc && install pandoc system-wide"
            )

    @staticmethod
    def _read_odt(file_path: str) -> str:
        """Read OpenDocument Text file."""
        if not ODT_AVAILABLE:
            raise ImportError("odfpy is required for ODT support. Install: pip install odfpy")

        doc = odf_load(file_path)
        paragraphs = doc.getElementsByType(text.P)
        text_content = '\n\n'.join(teletype.extractText(p) for p in paragraphs)
        return text_content

    @staticmethod
    def _read_latex(file_path: str) -> str:
        """Read LaTeX document (strips commands, keeps content)."""
        with open(file_path, 'r', encoding='utf-8') as f:
            content = f.read()

        # Remove comments
        content = re.sub(r'%.*$', '', content, flags=re.MULTILINE)

        # Remove common LaTeX commands while preserving content
        content = re.sub(r'\\documentclass\{.*?\}', '', content)
        content = re.sub(r'\\usepackage\{.*?\}', '', content)
        content = re.sub(r'\\begin\{document\}', '', content)
        content = re.sub(r'\\end\{document\}', '', content)
        content = re.sub(r'\\section\{(.*?)\}', r'\1\n', content)
        content = re.sub(r'\\subsection\{(.*?)\}', r'\1\n', content)
        content = re.sub(r'\\textbf\{(.*?)\}', r'\1', content)
        content = re.sub(r'\\textit\{(.*?)\}', r'\1', content)
        content = re.sub(r'\\cite\{.*?\}', '', content)
        content = re.sub(r'\\ref\{.*?\}', '', content)

        # Clean up whitespace
        content = re.sub(r'\n{3,}', '\n\n', content)

        return content.strip()

    # ========== OUTPUT WRITERS ==========

    @staticmethod
    def _write_txt(text: str, output_path: str, citation_style: str, metadata: Dict[str, Any]) -> None:
        """Write plain text file."""
        with open(output_path, 'w', encoding='utf-8') as f:
            if metadata.get('title'):
                f.write(f"{metadata['title']}\n")
                f.write("=" * len(metadata['title']) + "\n\n")
            if metadata.get('author'):
                f.write(f"Author: {metadata['author']}\n\n")
            f.write(text)

    @staticmethod
    def _write_markdown(text: str, output_path: str, citation_style: str, metadata: Dict[str, Any]) -> None:
        """Write Markdown file."""
        with open(output_path, 'w', encoding='utf-8') as f:
            # Add YAML front matter if metadata provided
            if metadata:
                f.write("---\n")
                if metadata.get('title'):
                    f.write(f"title: {metadata['title']}\n")
                if metadata.get('author'):
                    f.write(f"author: {metadata['author']}\n")
                if metadata.get('date'):
                    f.write(f"date: {metadata['date']}\n")
                f.write("---\n\n")

            f.write(text)

    @staticmethod
    def _write_docx(text: str, output_path: str, citation_style: str, metadata: Dict[str, Any]) -> None:
        """Write Microsoft Word document with academic formatting."""
        if not DOCX_AVAILABLE:
            raise ImportError("python-docx is required for DOCX output. Install: pip install python-docx")

        doc = Document()

        # Set margins (1 inch all around)
        sections = doc.sections
        for section in sections:
            section.top_margin = Inches(1)
            section.bottom_margin = Inches(1)
            section.left_margin = Inches(1)
            section.right_margin = Inches(1)

        # Add title
        if metadata.get('title'):
            title = doc.add_heading(metadata['title'], level=0)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Add author
        if metadata.get('author'):
            author = doc.add_paragraph(metadata['author'])
            author.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = author.runs[0]
            run.font.size = Pt(12)

        # Add date
        if metadata.get('date'):
            date_para = doc.add_paragraph(metadata['date'])
            date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = date_para.runs[0]
            run.font.size = Pt(12)

        if metadata:
            doc.add_paragraph()  # Spacing

        # Add content paragraphs
        paragraphs = text.split('\n\n')
        for para_text in paragraphs:
            if para_text.strip():
                p = doc.add_paragraph(para_text.strip())
                # Set font to Times New Roman, 12pt
                for run in p.runs:
                    run.font.name = 'Times New Roman'
                    run.font.size = Pt(12)
                # Double spacing
                p.paragraph_format.line_spacing = 2.0
                # First line indent (0.5 inch)
                p.paragraph_format.first_line_indent = Inches(0.5)

        doc.save(output_path)

    @staticmethod
    def _write_html(text: str, output_path: str, citation_style: str, metadata: Dict[str, Any]) -> None:
        """Write HTML file with academic styling."""
        if not HTML_AVAILABLE:
            raise ImportError("markdown is required for HTML output. Install: pip install markdown")

        # Convert to HTML
        html_content = markdown.markdown(text, extensions=['extra', 'nl2br'])

        # Create full HTML document
        html_doc = f"""<!DOCTYPE html>
<html lang="en">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>{metadata.get('title', 'Document')}</title>
    <style>
        body {{
            font-family: 'Times New Roman', Times, serif;
            font-size: 12pt;
            line-height: 2.0;
            max-width: 8.5in;
            margin: 1in auto;
            padding: 0 1in;
        }}
        h1 {{
            text-align: center;
            margin-bottom: 0.5em;
        }}
        .author, .date {{
            text-align: center;
            margin-bottom: 0.5em;
        }}
        p {{
            text-indent: 0.5in;
            margin-bottom: 0;
        }}
        @media print {{
            body {{
                margin: 1in;
            }}
        }}
    </style>
</head>
<body>
"""

        if metadata.get('title'):
            html_doc += f"    <h1>{metadata['title']}</h1>\n"
        if metadata.get('author'):
            html_doc += f"    <p class='author'>{metadata['author']}</p>\n"
        if metadata.get('date'):
            html_doc += f"    <p class='date'>{metadata['date']}</p>\n"

        html_doc += f"    {html_content}\n"
        html_doc += "</body>\n</html>"

        with open(output_path, 'w', encoding='utf-8') as f:
            f.write(html_doc)

    @staticmethod
    def _write_rtf(text: str, output_path: str, citation_style: str, metadata: Dict[str, Any]) -> None:
        """Write RTF file using pandoc."""
        if not PANDOC_AVAILABLE:
            raise ImportError(
                "pypandoc (with pandoc installed) is required for RTF output. "
                "Install: pip install pypandoc && install pandoc system-wide"
            )

        # Create markdown with metadata
        md_content = ""
        if metadata:
            md_content += "---\n"
            if metadata.get('title'):
                md_content += f"title: {metadata['title']}\n"
            if metadata.get('author'):
                md_content += f"author: {metadata['author']}\n"
            md_content += "---\n\n"
        md_content += text

        # Convert to RTF
        pypandoc.convert_text(md_content, 'rtf', format='md', outputfile=output_path)

    @staticmethod
    def _write_odt(text: str, output_path: str, citation_style: str, metadata: Dict[str, Any]) -> None:
        """Write OpenDocument Text file."""
        if not ODT_AVAILABLE:
            raise ImportError("odfpy is required for ODT output. Install: pip install odfpy")

        doc = OpenDocumentText()

        # Add title
        if metadata.get('title'):
            h = text.H(outlinelevel=1, text=metadata['title'])
            doc.text.addElement(h)

        # Add author
        if metadata.get('author'):
            p = text.P(text=metadata['author'])
            doc.text.addElement(p)

        # Add paragraphs
        paragraphs = text.split('\n\n')
        for para_text in paragraphs:
            if para_text.strip():
                p = text.P(text=para_text.strip())
                doc.text.addElement(p)

        doc.save(output_path)

    @staticmethod
    def _write_latex(text: str, output_path: str, citation_style: str, metadata: Dict[str, Any]) -> None:
        """Write LaTeX document."""
        latex_doc = r"""\documentclass[12pt]{article}
\usepackage[margin=1in]{geometry}
\usepackage{setspace}
\doublespacing

"""

        # Add title/author if provided
        if metadata.get('title') or metadata.get('author'):
            if metadata.get('title'):
                latex_doc += f"\\title{{{metadata['title']}}}\n"
            if metadata.get('author'):
                latex_doc += f"\\author{{{metadata['author']}}}\n"
            if metadata.get('date'):
                latex_doc += f"\\date{{{metadata['date']}}}\n"
            else:
                latex_doc += "\\date{}\n"

        latex_doc += "\n\\begin{document}\n"

        if metadata.get('title'):
            latex_doc += "\\maketitle\n\n"

        # Escape special LaTeX characters
        text_escaped = text.replace('&', '\\&')
        text_escaped = text_escaped.replace('%', '\\%')
        text_escaped = text_escaped.replace('$', '\\$')
        text_escaped = text_escaped.replace('#', '\\#')
        text_escaped = text_escaped.replace('_', '\\_')
        text_escaped = text_escaped.replace('{', '\\{')
        text_escaped = text_escaped.replace('}', '\\}')
        text_escaped = text_escaped.replace('~', '\\~{}')
        text_escaped = text_escaped.replace('^', '\\^{}')

        # Add paragraphs
        paragraphs = text_escaped.split('\n\n')
        for para in paragraphs:
            if para.strip():
                latex_doc += f"{para.strip()}\n\n"

        latex_doc += "\\end{document}\n"

        with open(output_path, 'w', encoding='utf-8') as f:
            f.write(latex_doc)

    @staticmethod
    def _write_pdf(text: str, output_path: str, citation_style: str, metadata: Dict[str, Any]) -> None:
        """Write PDF file (using weasyprint if available, else reportlab)."""
        if WEASYPRINT_AVAILABLE:
            # Use weasyprint (better quality, CSS support)
            html_content = f"""
            <html>
            <head>
                <style>
                    @page {{ margin: 1in; }}
                    body {{
                        font-family: 'Times New Roman', Times, serif;
                        font-size: 12pt;
                        line-height: 2.0;
                    }}
                    h1 {{ text-align: center; }}
                    .author, .date {{ text-align: center; }}
                    p {{ text-indent: 0.5in; margin-bottom: 0; }}
                </style>
            </head>
            <body>
            """
            if metadata.get('title'):
                html_content += f"<h1>{metadata['title']}</h1>"
            if metadata.get('author'):
                html_content += f"<p class='author'>{metadata['author']}</p>"
            if metadata.get('date'):
                html_content += f"<p class='date'>{metadata['date']}</p>"

            paragraphs = text.split('\n\n')
            for para in paragraphs:
                if para.strip():
                    html_content += f"<p>{para.strip()}</p>"

            html_content += "</body></html>"

            WeasyHTML(string=html_content).write_pdf(output_path)

        elif REPORTLAB_AVAILABLE:
            # Use reportlab (fallback)
            doc = SimpleDocTemplate(output_path, pagesize=letter)
            styles = getSampleStyleSheet()
            story = []

            if metadata.get('title'):
                story.append(Paragraph(metadata['title'], styles['Title']))
                story.append(Spacer(1, 12))

            paragraphs = text.split('\n\n')
            for para in paragraphs:
                if para.strip():
                    story.append(Paragraph(para.strip(), styles['Normal']))
                    story.append(Spacer(1, 12))

            doc.build(story)
        else:
            raise ImportError(
                "weasyprint or reportlab is required for PDF output. "
                "Install: pip install weasyprint OR pip install reportlab"
            )


# Convenience functions
def read_document(file_path: str, format: Optional[DocumentFormat] = None) -> str:
    """Convenience function to read a document.

    Args:
        file_path: Path to the document
        format: Document format (auto-detected if None)

    Returns:
        Document text content
    """
    return DocumentFormatHandler.read_document(file_path, format)


def write_document(
    text: str,
    output_path: str,
    format: Optional[DocumentFormat] = None,
    citation_style: Literal["apa", "chicago"] = "apa",
    metadata: Optional[Dict[str, Any]] = None
) -> None:
    """Convenience function to write a document.

    Args:
        text: Document text content (with citations)
        output_path: Path to save the document
        format: Document format (auto-detected if None)
        citation_style: Citation style for formatting
        metadata: Optional metadata (title, author, etc.)
    """
    DocumentFormatHandler.write_document(text, output_path, format, citation_style, metadata)
