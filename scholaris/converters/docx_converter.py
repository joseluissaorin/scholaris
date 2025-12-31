"""Markdown to DOCX converter for literature reviews.

Adapted from original docx_converter.py to work with Scholaris Review objects.
"""
import os
import tempfile
import shutil
import urllib.request
import logging
from typing import Optional

import markdown
from bs4 import BeautifulSoup

try:
    from docx import Document
    from docx.shared import Pt, Cm, Inches, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
    from docx.enum.section import WD_ORIENT
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    import docx.opc.constants
    import docx.enum.style
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

from ..exceptions import ConversionError

logger = logging.getLogger(__name__)


class DocxConverter:
    """Convert markdown to DOCX format with academic styling."""

    def __init__(self, output_folder: Optional[str] = None):
        """Initialize DOCX converter.

        Args:
            output_folder: Optional output folder (uses temp if not provided)

        Raises:
            ConversionError: If python-docx not installed
        """
        if not DOCX_AVAILABLE:
            raise ConversionError(
                "python-docx not installed. Install it with: pip install python-docx"
            )

        self.output_folder = output_folder or tempfile.gettempdir()
        os.makedirs(self.output_folder, exist_ok=True)
        self.document = None
        self.temp_dir = tempfile.mkdtemp()
        self.footnotes = []

    def __del__(self):
        """Clean up temporary files."""
        try:
            if hasattr(self, 'temp_dir') and os.path.exists(self.temp_dir):
                shutil.rmtree(self.temp_dir)
        except Exception:
            pass

    def convert(self, markdown_text: str, output_path: str) -> str:
        """Convert markdown text to DOCX file.

        Args:
            markdown_text: Markdown content
            output_path: Path to save DOCX file

        Returns:
            Path to created DOCX file

        Raises:
            ConversionError: If conversion fails
        """
        logger.info(f"Converting markdown to DOCX: {output_path}")

        try:
            self.document = Document()
            self._set_document_styles()
            self.footnotes = []

            # Convert markdown to HTML
            html = markdown.markdown(markdown_text, extensions=[
                'extra',
                'tables',
                'sane_lists',
                'footnotes',
                'smarty',
                'codehilite',
                'nl2br'
            ])

            # Convert HTML to DOCX
            self._html_to_docx(html)

            # Add footnotes if any
            if self.footnotes:
                self._add_footnotes_section()

            # Save document
            self.document.save(output_path)
            logger.info(f"✓ DOCX saved: {output_path}")

            return output_path

        except Exception as e:
            logger.error(f"DOCX conversion failed: {e}")
            raise ConversionError(f"Failed to convert to DOCX: {e}")

    def _set_document_styles(self):
        """Configure document styles for academic format."""
        # Page layout - A4 with standard margins
        section = self.document.sections[0]
        section.page_height = Cm(29.7)
        section.page_width = Cm(21.0)
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)
        section.orientation = WD_ORIENT.PORTRAIT
        section.different_first_page_header_footer = True

        # Base font style
        style = self.document.styles['Normal']
        font = style.font
        font.name = 'Times New Roman'
        font.size = Pt(12)

        # Set font fallbacks
        rFonts = font._element.rPr.get_or_add_rFonts()
        rFonts.set(qn('w:ascii'), 'Times New Roman')
        rFonts.set(qn('w:hAnsi'), 'Times New Roman')
        rFonts.set(qn('w:cs'), 'Arial')
        rFonts.set(qn('w:eastAsia'), 'Arial')

        # Heading styles
        for i in range(1, 10):
            style_name = f'Heading {i}'
            if style_name in self.document.styles:
                heading_style = self.document.styles[style_name]
                heading_font = heading_style.font
                heading_font.name = 'Times New Roman'
                rFonts = heading_font._element.rPr.get_or_add_rFonts()
                rFonts.set(qn('w:ascii'), 'Times New Roman')
                rFonts.set(qn('w:hAnsi'), 'Times New Roman')

        # Code block style
        if 'Code Block' not in self.document.styles:
            code_style = self.document.styles.add_style(
                'Code Block',
                docx.enum.style.WD_STYLE_TYPE.PARAGRAPH
            )
            code_font = code_style.font
            code_font.name = 'Courier New'
            code_font.size = Pt(10)
            paragraph_format = code_style.paragraph_format
            paragraph_format.space_before = Pt(6)
            paragraph_format.space_after = Pt(6)
            paragraph_format.left_indent = Cm(1)

    def _html_to_docx(self, html: str):
        """Parse HTML and add to DOCX document."""
        soup = BeautifulSoup(html, 'html.parser')
        for element in soup.find_all(recursive=False):
            self._process_element(element)

    def _process_element(self, element):
        """Process an HTML element and add to document."""
        if element.name in ['p', 'li']:
            if element.name == 'li':
                list_style = 'List Bullet' if element.find_parent('ul') else 'List Number'
                paragraph = self.document.add_paragraph(style=list_style)
            else:
                paragraph = self.document.add_paragraph()
                paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                paragraph.paragraph_format.first_line_indent = Cm(1.27)
            self._process_inline_elements(element, paragraph)

        elif element.name in ['h1', 'h2', 'h3', 'h4', 'h5', 'h6']:
            level = int(element.name[1])
            heading = self.document.add_heading(level=level)
            self._process_inline_elements(element, heading)

        elif element.name == 'blockquote':
            paragraph = self.document.add_paragraph(style='Quote')
            for child in element.children:
                self._process_element(child)

        elif element.name == 'pre':
            code_element = element.find('code')
            if code_element:
                paragraph = self.document.add_paragraph(style='Code Block')
                run = paragraph.add_run(code_element.get_text())
                run.font.name = 'Courier New'
                run.font.size = Pt(10)

        elif element.name == 'img':
            self._add_image(element)

        elif element.name == 'table':
            self._add_table(element)

        elif element.name == 'div' and 'footnote' in element.get('class', []):
            footnote_id = element.get('id', '').replace('fn:', '')
            footnote_content = ' '.join(element.get_text().strip().split())
            if footnote_id and footnote_content:
                self.footnotes.append((footnote_id, footnote_content))

        else:
            for child in element.children:
                if isinstance(child, str):
                    if child.strip():
                        p = self.document.add_paragraph()
                        p.add_run(child)
                else:
                    self._process_element(child)

    def _add_image(self, img_element):
        """Add image to document."""
        src = img_element.get('src', '')
        if not src:
            return

        try:
            if src.startswith(('http://', 'https://')):
                img_temp_path = os.path.join(
                    self.temp_dir,
                    f"img_{hash(src)}.png"
                )
                urllib.request.urlretrieve(src, img_temp_path)
                img_path = img_temp_path
            else:
                if not os.path.exists(src):
                    return
                img_path = src

            # Add image with reasonable size
            width = img_element.get('width')
            if width and width.isdigit():
                width_inches = min(6, int(width) / 96)
                self.document.add_picture(img_path, width=Inches(width_inches))
            else:
                self.document.add_picture(img_path, width=Inches(5))

            # Add caption if alt text present
            alt_text = img_element.get('alt', '')
            if alt_text:
                caption = self.document.add_paragraph(alt_text, style='Caption')
                caption.alignment = WD_ALIGN_PARAGRAPH.CENTER

        except Exception as e:
            logger.warning(f"Failed to add image {src}: {e}")
            p = self.document.add_paragraph(f"[Image: {src}]", style='Caption')
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    def _add_table(self, table_element):
        """Add table to document."""
        rows = table_element.find_all('tr')
        if not rows:
            return

        max_cells = max(len(row.find_all(['td', 'th'])) for row in rows)
        if max_cells == 0:
            return

        table = self.document.add_table(rows=len(rows), cols=max_cells)
        table.style = 'Table Grid'

        for i, row in enumerate(rows):
            cells = row.find_all(['td', 'th'])
            has_header = cells and cells[0].name == 'th'

            for j, cell in enumerate(cells):
                if j < max_cells:
                    table_cell = table.cell(i, j)
                    p = table_cell.paragraphs[0] if table_cell.paragraphs else table_cell.add_paragraph()
                    self._process_inline_elements(cell, p)

                    if has_header or cell.name == 'th':
                        for run in p.runs:
                            run.bold = True

        self.document.add_paragraph()

    def _add_footnotes_section(self):
        """Add footnotes section at end of document."""
        if not self.footnotes:
            return

        para = self.document.add_paragraph()
        run = para.add_run()
        run.add_break(WD_BREAK.PAGE)

        self.document.add_heading('Footnotes', level=1)

        for footnote_id, content in self.footnotes:
            footnote_para = self.document.add_paragraph()
            footnote_para.add_run(f"{footnote_id}. ").bold = True
            footnote_para.add_run(content)

    def _process_inline_elements(self, element, paragraph):
        """Process inline elements within a paragraph."""
        for node in element.contents:
            if node.name is None:
                text = str(node)
                if text.strip():
                    paragraph.add_run(text)
            else:
                text = node.get_text()
                if node.name in ['strong', 'b']:
                    run = paragraph.add_run(text)
                    run.bold = True
                elif node.name in ['em', 'i']:
                    run = paragraph.add_run(text)
                    run.italic = True
                elif node.name == 'u':
                    run = paragraph.add_run(text)
                    run.underline = True
                elif node.name == 'code':
                    run = paragraph.add_run(text)
                    run.font.name = 'Courier New'
                    run.font.size = Pt(10)
                elif node.name == 'a':
                    href = node.get('href', '')
                    if href:
                        run = paragraph.add_run(text)
                        run.underline = True
                        run.font.color.rgb = RGBColor(0, 0, 255)
                        self._add_hyperlink(paragraph, href, text)

    def _add_hyperlink(self, paragraph, url, text):
        """Add hyperlink to paragraph."""
        try:
            part = paragraph.part
            r_id = part.relate_to(
                url,
                docx.opc.constants.RELATIONSHIP_TYPE.HYPERLINK,
                is_external=True
            )
            return r_id
        except Exception:
            return None
